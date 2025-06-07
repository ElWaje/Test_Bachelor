import os
import sys
import json
import shutil
import random
import re
from glob import glob
from docx import Document
from PyQt5 import QtWidgets, QtCore
from PyQt5.QtGui import QPainter, QPixmap
from collections import defaultdict
from datetime import datetime
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
import matplotlib.pyplot as plt
import numpy as np
import csv

# Constants
tests_base = './tests_json'
results_file = 'results.json'

def ensure_base():
    os.makedirs(tests_base, exist_ok=True)

def load_sections():
    ensure_base()
    secs = [d for d in os.listdir(tests_base)
            if os.path.isdir(os.path.join(tests_base, d))]
    if not secs:
        jsons = glob(os.path.join(tests_base, '*.json'))
        if jsons:
            secs = ['General']
    return secs

def get_subject_and_topic(path):
    base = os.path.basename(path)
    match = re.match(r'(.+?)_+Tema(\d+)', base, re.IGNORECASE)
    if match:
        subject = match.group(1).replace('_', ' ')
        topic = int(match.group(2))
        return subject.strip(), topic
    return None, None

def select_half_questions(half='first', section='General'):
    all_files = []
    pattern = os.path.join(tests_base, '*.json') if section == 'General' else os.path.join(tests_base, section, '*.json')
    for f in glob(pattern):
        all_files.append((f, section))
    grouped = {}
    for f, sec in all_files:
        subject, topic = get_subject_and_topic(f)
        if subject is None or topic is None:
            continue
        grouped.setdefault(subject, []).append((topic, f, sec))
    selected_questions = []
    for subject, lst in grouped.items():
        lst.sort()
        mid = len(lst) // 2
        subset = lst[:mid] if half == 'first' else lst[mid:]
        for _, fpath, sec in subset:
            questions = parse_test(fpath)
            for q in questions:
                q['section'] = sec
                q['file'] = fpath
                selected_questions.append(q)
    return random.sample(selected_questions, min(40, len(selected_questions)))

def parse_test(path):
    if path.lower().endswith('.json'):
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    doc = Document(path)
    qlist, current = [], None
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue
        if re.match(r'.+\?$', t) and not re.match(r'^[ABCD][\.)]', t):
            if current: qlist.append(current)
            current = {'question': t, 'options': {}, 'answer': None}
            continue
        if current and re.match(r'^[ABCD][\.\)\-]\s+', t):
            label, text = t[0], re.sub(r'^[ABCD][\.\)\-]\s+', '', t)
            current['options'][label] = text
            continue
        if current and re.match(r'(?i)^respuesta[:]?[ \t]*[ABCD]', t):
            m = re.search(r'([ABCD])', t.upper())
            if m: current['answer'] = m.group(1)
            continue
    if current: qlist.append(current)
    if any(q.get('answer') is None for q in qlist):
        answers = [p.text.strip()[0] for p in doc.paragraphs
                   if re.match(r'^[ABCD][\.\)\-]', p.text.strip())]
        if len(answers) >= len(qlist):
            for q, a in zip(qlist, answers): q['answer'] = a
    return qlist

def load_results():
    if os.path.exists(results_file):
        with open(results_file,'r',encoding='utf-8') as f:
            return json.load(f)
    return {'exams':[]}

def save_results(data):
    with open(results_file,'w',encoding='utf-8') as f:
        json.dump(data,f,indent=2,ensure_ascii=False)

def show_review(parent, answers):
    dlg = QtWidgets.QDialog(parent)
    dlg.setWindowTitle('Revisión del Test')
    dlg.resize(600, 400)
    lay = QtWidgets.QVBoxLayout(dlg)
    html = '<html><body>'
    for item in answers:
        q = item['question']
        corr_txt = item.get('correct_text', '–––')
        sel_txt = item.get('selected_text', 'Sin responder')
        corr_letter = item.get('correct_letter')
        sel_letter = item.get('selected_letter')
        color_sel = 'green' if sel_letter == corr_letter else 'red'
        html += f'<p><b>Pregunta:</b> {q}<br>'
        html += f'<span style="color:green"><b>Correcta:</b> {corr_txt}</span><br>'
        html += f'<span style="color:{color_sel}"><b>Tu respuesta:</b> {sel_txt}</span></p><hr>'
    html += '</body></html>'
    te = QtWidgets.QTextEdit()
    te.setReadOnly(True)
    te.setHtml(html)
    lay.addWidget(te)
    btn = QtWidgets.QPushButton('Cerrar')
    btn.clicked.connect(dlg.accept)
    lay.addWidget(btn)
    dlg.exec_()

# NUEVO: Widget con fondo de imagen
class FondoWidget(QtWidgets.QWidget):
    def __init__(self, image_path, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._image = QPixmap(image_path)

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.drawPixmap(self.rect(), self._image)
        super().paintEvent(event)

class TestApp(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('Generador de Tests')
        self.resize(1000,700)
        self.results = load_results()
        self._init_ui()

    def _init_ui(self):
        tabs = QtWidgets.QTabWidget()
        tabs.addTab(self._tests_tab(),'Tests')
        tabs.addTab(self._stats_tab(),'Estadísticas')
        self.setCentralWidget(tabs)

    def _tests_tab(self):
        # Usa FondoWidget en vez de QWidget
        w = FondoWidget('fondo.png')  # Debe estar en la misma carpeta
        layout = QtWidgets.QVBoxLayout()
        hl = QtWidgets.QHBoxLayout()
        self.section_cb = QtWidgets.QComboBox(); hl.addWidget(self.section_cb)
        self.test_cb = QtWidgets.QComboBox(); hl.addWidget(self.test_cb)
        btn_load = QtWidgets.QPushButton('Hacer Test'); btn_load.clicked.connect(self.start_lesson)
        hl.addWidget(btn_load)
        layout.addLayout(hl)
        hl2 = QtWidgets.QHBoxLayout()
        btn_rand = QtWidgets.QPushButton('Examen 20 aleatorias'); btn_rand.clicked.connect(self.start_random)
        hl2.addWidget(btn_rand)
        btn_add_sec = QtWidgets.QPushButton('Nueva Sección'); btn_add_sec.clicked.connect(self.add_section)
        hl2.addWidget(btn_add_sec)
        btn_add_test = QtWidgets.QPushButton('Añadir Test'); btn_add_test.clicked.connect(self.add_test)
        hl2.addWidget(btn_add_test)
        btn_first_half = QtWidgets.QPushButton('Examen 40 (1ª mitad temas)')
        btn_first_half.clicked.connect(lambda: self.start_half_test('first'))
        hl2.addWidget(btn_first_half)
        btn_second_half = QtWidgets.QPushButton('Examen 40 (2ª mitad temas)')
        btn_second_half.clicked.connect(lambda: self.start_half_test('second'))
        hl2.addWidget(btn_second_half)
        layout.addLayout(hl2)
        self._refresh_sections()
        w.setLayout(layout)
        return w

    def _stats_tab(self):
        w = QtWidgets.QWidget()
        layout = QtWidgets.QVBoxLayout()
        self.stats_txt = QtWidgets.QTextEdit(); self.stats_txt.setReadOnly(True)
        layout.addWidget(self.stats_txt)
        self.stats_table = QtWidgets.QTableWidget()
        self.stats_table.setColumnCount(6)
        self.stats_table.setHorizontalHeaderLabels(['Fecha', 'Sección', 'Tipo', 'Test', 'Aciertos/Total', 'Porcentaje'])
        layout.addWidget(self.stats_table)
        self.figure1, self.ax1 = plt.subplots(figsize=(5,2))
        self.canvas1 = FigureCanvas(self.figure1)
        layout.addWidget(self.canvas1)
        self.figure2, self.ax2 = plt.subplots(figsize=(5,2))
        self.canvas2 = FigureCanvas(self.figure2)
        layout.addWidget(self.canvas2)
        btn_layout = QtWidgets.QHBoxLayout()
        btn = QtWidgets.QPushButton('Actualizar')
        btn.clicked.connect(self.show_stats)
        btn_layout.addWidget(btn)
        btn_export = QtWidgets.QPushButton('Exportar CSV')
        btn_export.clicked.connect(self.export_csv)
        btn_layout.addWidget(btn_export)
        layout.addLayout(btn_layout)
        w.setLayout(layout)
        return w

    def export_csv(self):
        exams = self.results.get('exams', [])
        if not exams:
            QtWidgets.QMessageBox.warning(self, 'Sin datos', 'No hay exámenes para exportar.')
            return
        fname, _ = QtWidgets.QFileDialog.getSaveFileName(self, 'Guardar CSV', '', 'CSV Files (*.csv)')
        if fname:
            with open(fname, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow(['Fecha', 'Sección', 'Tipo', 'Test', 'Aciertos', 'Total', 'Porcentaje'])
                for e in exams:
                    fecha = e.get('date', '')[:19].replace('T', ' ')
                    sec = e.get('section', 'General')
                    typ = e.get('type', '')
                    test = e.get('test', '')
                    score = e.get('score', 0)
                    total = e.get('total', 0)
                    pct = f"{(score/total*100):.2f}%" if total else "-"
                    writer.writerow([fecha, sec, typ, test, score, total, pct])
            QtWidgets.QMessageBox.information(self, 'CSV exportado', 'Exportación completada.')

    def _refresh_sections(self):
        secs = load_sections()
        self.section_cb.blockSignals(True)
        self.section_cb.clear(); self.section_cb.addItems(secs)
        self.section_cb.blockSignals(False)
        self._refresh_tests()
        self.section_cb.currentIndexChanged.connect(self._refresh_tests)

    def _refresh_tests(self):
        sec = self.section_cb.currentText()
        pattern = os.path.join(tests_base, '*.json') if sec == 'General' else os.path.join(tests_base, sec, '*.json')
        files = [os.path.basename(f) for f in glob(pattern)]
        self.test_cb.clear(); self.test_cb.addItems(files)

    def add_section(self):
        name, ok = QtWidgets.QInputDialog.getText(self, 'Nueva Sección', 'Nombre:')
        if ok and name:
            sec = name.replace(' ', '_')
            os.makedirs(os.path.join(tests_base, sec), exist_ok=True)
            self._refresh_sections()

    def add_test(self):
        sec = self.section_cb.currentText()
        src, _ = QtWidgets.QFileDialog.getOpenFileName(self, 'Seleccionar .json', '', 'JSON files (*.json)')
        if src and src.lower().endswith('.json'):
            dst = os.path.join(tests_base, os.path.basename(src)) if sec == 'General' else os.path.join(tests_base, sec, os.path.basename(src))
            shutil.copy(src, dst); self._refresh_tests()

    def start_lesson(self):
        sec = self.section_cb.currentText(); test = self.test_cb.currentText()
        if not test:
            QtWidgets.QMessageBox.warning(self, 'Error', 'No hay tests en la sección seleccionada.')
            return
        path = os.path.join(tests_base, test) if sec == 'General' else os.path.join(tests_base, sec, test)
        qs = parse_test(path)
        if not qs:
            QtWidgets.QMessageBox.warning(self, 'Error', f'No se encontraron preguntas en {test}.')
            return
        self._conduct(qs, 'lesson', sec, test)

    def start_random(self):
        sec = self.section_cb.currentText()
        pattern = os.path.join(tests_base, '*.json') if sec == 'General' else os.path.join(tests_base, sec, '*.json')
        all_q = []
        for f in glob(pattern):
            for q in parse_test(f):
                q['section'] = sec
                all_q.append(q)
        if not all_q:
            QtWidgets.QMessageBox.warning(self, 'Error', 'No hay preguntas disponibles en la sección seleccionada.')
            return
        sel = random.sample(all_q, min(20, len(all_q)))
        self._conduct(sel, 'random20', sec, None)

    def start_half_test(self, half):
        sec = self.section_cb.currentText()
        sel = select_half_questions(half, sec)
        if not sel:
            QtWidgets.QMessageBox.warning(self, 'Error', f'No hay suficientes preguntas en la {half} mitad de los temas de la sección seleccionada.')
            return
        self._conduct(sel, f'half_{half}', sec, None)

    def _conduct(self, questions, typ, sec, test):
        dlg = QuestionDialog(questions, self)
        if dlg.exec_():
            ans = dlg.user_answers
            score = sum(1 for a in ans if a.get('selected_letter') == a.get('correct_letter'))
            self.results['exams'].append({
                'type': typ,
                'section': sec,
                'test': test,
                'score': score,
                'total': len(ans),
                'date': datetime.now().isoformat(timespec='seconds')
            })
            save_results(self.results)
            QtWidgets.QMessageBox.information(self, 'Resultados', f'Puntuación: {score}/{len(ans)}')
            self.show_stats()
            show_review(self, ans)

    def show_stats(self):
        exams = self.results.get('exams', [])
        # Resumen clásico
        by_sec = {}
        total_c = total_q = 0
        for e in exams:
            key = e.get('section', 'General')
            by_sec.setdefault(key, []).append(e)
            total_c += e['score']
            total_q += e['total']
        text = f"--- Estadísticas globales ---\n"
        text += f"Exámenes realizados: {len(exams)}\n"
        text += f"Puntuación total acumulada: {total_c}/{total_q}\n"
        media_global = (total_c / total_q * 100) if total_q else 0
        text += f"Puntuación media global: {media_global:.2f}%\n"

        # Por fecha
        by_date = defaultdict(list)
        for e in exams:
            fecha = e.get('date', '')[:10]
            by_date[fecha].append(e)
        text += "\n--- Por fecha ---\n"
        for fecha in sorted(by_date):
            exs = by_date[fecha]
            sc = sum(e['score'] for e in exs)
            tq = sum(e['total'] for e in exs)
            med = (sc / tq * 100) if tq else 0
            text += f"{fecha}: {len(exs)} exámenes, media {med:.2f}%\n"

        # Resumen por tipo de test
        by_type = defaultdict(list)
        for e in exams:
            by_type[e.get('type', 'Desconocido')].append(e)
        text += "\n--- Por tipo de test ---\n"
        for typ, lst in by_type.items():
            score = sum(e['score'] for e in lst)
            total = sum(e['total'] for e in lst)
            media = (score/total*100) if total else 0
            text += f"{typ}: {len(lst)} exámenes, media {media:.2f}%\n"

        self.stats_txt.setText(text)

        # Tabla con color
        self.stats_table.setRowCount(len(exams))
        for row, e in enumerate(sorted(exams, key=lambda x: x.get('date', ''))):
            fecha = e.get('date', '')[:19].replace('T', ' ')
            sec = str(e.get('section', 'General'))
            typ = str(e.get('type', ''))
            test = str(e.get('test', ''))
            score = f"{e['score']}/{e['total']}"
            percent_val = (e['score']/e['total']*100) if e['total'] else 0
            percent = f"{percent_val:.2f}%"
            for col, val in enumerate([fecha, sec, typ, test, score, percent]):
                item = QtWidgets.QTableWidgetItem(val)
                if col == 5:
                    # Colores: verde >=80, amarillo >=60, rojo <60
                    if percent_val >= 80:
                        item.setBackground(QtCore.Qt.green)
                    elif percent_val >= 60:
                        item.setBackground(QtCore.Qt.yellow)
                    else:
                        item.setBackground(QtCore.Qt.red)
                    item.setTextAlignment(QtCore.Qt.AlignCenter)
                self.stats_table.setItem(row, col, item)
        self.stats_table.resizeColumnsToContents()

        # Gráfica evolución diaria con línea de tendencia
        self.ax1.clear()
        fechas = sorted(by_date)
        x = np.arange(len(fechas))
        medias = [(sum(e['score'] for e in by_date[f]) / sum(e['total'] for e in by_date[f]) * 100)
                  if sum(e['total'] for e in by_date[f]) else 0 for f in fechas]
        self.ax1.plot(fechas, medias, marker='o', linestyle='-', color='tab:blue', label='Media diaria (%)')
        # Línea de tendencia polinómica grado 1 (recta)
        if len(medias) > 1:
            z = np.polyfit(x, medias, 1)
            p = np.poly1d(z)
            self.ax1.plot(fechas, p(x), color='tab:red', linestyle='--', label='Tendencia')
        self.ax1.set_title("Evolución diaria de la puntuación")
        self.ax1.set_xlabel("Fecha")
        self.ax1.set_ylabel("Media (%)")
        self.ax1.set_ylim(0, 105)
        self.ax1.grid(True)
        self.ax1.legend()
        self.figure1.tight_layout()
        self.canvas1.draw()

        # Gráfica de barras por tipo de test
        self.ax2.clear()
        tipos = list(by_type.keys())
        medias_tipo = [(sum(e['score'] for e in by_type[t]) / sum(e['total'] for e in by_type[t]) * 100)
                       if sum(e['total'] for e in by_type[t]) else 0 for t in tipos]
        self.ax2.bar(tipos, medias_tipo, color='tab:green', alpha=0.7)
        self.ax2.set_title('Media por tipo de test')
        self.ax2.set_xlabel('Tipo de test')
        self.ax2.set_ylabel('Media (%)')
        self.ax2.set_ylim(0, 105)
        self.ax2.grid(axis='y')
        self.figure2.tight_layout()
        self.canvas2.draw()

class QuestionDialog(QtWidgets.QDialog):
    def __init__(self, questions, parent=None):
        super().__init__(parent)
        self.questions = questions; self.user_answers = []; self.idx = 0
        self.setWindowTitle('Examen')
        self.layout = QtWidgets.QVBoxLayout(); self.setLayout(self.layout)
        self.lbl = QtWidgets.QLabel(); self.lbl.setWordWrap(True)
        self.layout.addWidget(self.lbl)
        self.opt_group = QtWidgets.QGroupBox('Opciones')
        self.opt_layout = QtWidgets.QVBoxLayout(); self.opt_group.setLayout(self.opt_layout)
        self.layout.addWidget(self.opt_group)
        self.radio_group = QtWidgets.QButtonGroup(self); self.radio_group.setExclusive(True)
        self.btn_next = QtWidgets.QPushButton('Siguiente'); self.btn_next.clicked.connect(self.next_q)
        self.layout.addWidget(self.btn_next)
        self._show_question()

    def _show_question(self):
        q = self.questions[self.idx]; self.lbl.setText(q['question'])
        for i in reversed(range(self.opt_layout.count())):
            w = self.opt_layout.itemAt(i).widget(); self.radio_group.removeButton(w)
            self.opt_layout.removeWidget(w); w.deleteLater()
        for key, txt in sorted(q['options'].items()):
            rb = QtWidgets.QRadioButton(f"{key}. {txt}")
            rb.setAutoExclusive(False); rb.setChecked(False)
            self.radio_group.addButton(rb); self.opt_layout.addWidget(rb)
        for b in self.radio_group.buttons():
            b.toggled.connect(lambda checked, btn=b: btn.setAutoExclusive(True) if checked else None)
        if self.idx == len(self.questions) - 1:
            self.btn_next.setText('Terminar')

    def next_q(self):
        sel_letter = None
        for b in self.radio_group.buttons():
            if b.isChecked():
                sel_letter = b.text()[0]
        corr_letter = self.questions[self.idx]['answer']
        sel_text = None
        if sel_letter:
            sel_text = self.questions[self.idx]['options'].get(sel_letter, None)
        corr_text = self.questions[self.idx]['options'].get(corr_letter, None)
        self.user_answers.append({
            'question': self.questions[self.idx]['question'],
            'selected_letter': sel_letter,
            'selected_text': sel_text,
            'correct_letter': corr_letter,
            'correct_text': corr_text
        })
        self.idx += 1
        if self.idx < len(self.questions):
            self._show_question()
        else:
            self.accept()

if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    win = TestApp(); win.show(); sys.exit(app.exec_())
